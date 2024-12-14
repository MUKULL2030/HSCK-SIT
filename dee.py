import argparse
import PyPDF2
import docx
import nltk
from nltk.corpus import stopwords
from nltk.tokenize import word_tokenize
import torch
from transformers import BertTokenizer, BertModel, pipeline

nltk.download('punkt')
nltk.download('stopwords')

# Load BERT model and tokenizer
tokenizer = BertTokenizer.from_pretrained('bert-base-uncased')
model = BertModel.from_pretrained('bert-base-uncased')

# Load zero-shot classification model
classifier = pipeline("zero-shot-classification", model="facebook/bart-large-mnli")

def validate_file_format(file_path):
    if file_path.endswith('.pdf'):
        try:
            with open(file_path, 'rb') as f:
                PyPDF2.PdfReader(f)
            return True
        except PyPDF2.errors.PdfReadError:
            return False
    elif file_path.endswith('.docx'):
        try:
            docx.Document(file_path)
            return True
        except docx.exceptions.DocxPackageNotFoundError:
            return False
    else:
        return False

def extract_text_from_pdf(file_path):
    with open(file_path, 'rb') as f:
        pdf_reader = PyPDF2.PdfReader(f)
        text = ''
        for page in range(len(pdf_reader.pages)):
            text += pdf_reader.pages[page].extract_text()
        return text

def extract_text_from_docx(file_path):
    doc = docx.Document(file_path)
    text = ''
    for para in doc.paragraphs:
        text += para.text
    return text

def extract_keywords(text):
    stop_words = set(nltk.corpus.stopwords.words('english'))
    word_tokens = nltk.word_tokenize(text)
    keywords = [word for word in word_tokens if word not in stop_words]
    return keywords

def parse_resume(text):
    keywords = extract_keywords(text)
    
    # Placeholder for a more sophisticated extraction logic
    skills = [word for word in keywords if word.lower() in ['python', 'machine learning', 'nlp', 'data analysis']]
    experience_years = 0
    if 'years' in keywords:
        years_index = keywords.index('years')
        if years_index > 0 and keywords[years_index - 1].isdigit():
            experience_years = int(keywords[years_index - 1])
    projects = ['project' for word in keywords if 'project' in word.lower()]
    
    return skills, experience_years, projects

def generate_insights(skills, experience_years, projects):
    # Combine all information into one text
    text = f"Skills: {', '.join(skills)}. Experience Years: {experience_years}. Projects: {', '.join(projects)}."
    
    # Define the labels (categories) for classification
    labels = ["strengths", "weaknesses", "insights", "suggestions"]
    
    # Get the classification results
    classification_result = classifier(text, candidate_labels=labels)
    
    # Extract the top results for each label
    strengths = "No strengths identified"
    weaknesses = "No weaknesses identified"
    insights = "No insights identified"
    suggestions = "No suggestions identified"

    if 'strengths' in classification_result['labels']:
        strengths = classification_result['labels'][classification_result['labels'].index('strengths')]
    if 'weaknesses' in classification_result['labels']:
        weaknesses = classification_result['labels'][classification_result['labels'].index('weaknesses')]
    if 'insights' in classification_result['labels']:
        insights = classification_result['labels'][classification_result['labels'].index('insights')]
    if 'suggestions' in classification_result['labels']:
        suggestions = classification_result['labels'][classification_result['labels'].index('suggestions')]
    
    return strengths, weaknesses, insights, suggestions

def main(file_path):
    if validate_file_format(file_path):
        if file_path.endswith('.pdf'):
            text = extract_text_from_pdf(file_path)
        elif file_path.endswith('.docx'):
            text = extract_text_from_docx(file_path)

        # Extracting skills, experience_years, and projects
        skills, experience_years, projects = parse_resume(text)

        # Generate insights using BERT
        strengths, weaknesses, insights, suggestions = generate_insights(skills, experience_years, projects)

        # Print the results
        print(f"Strengths: {strengths}")
        print(f"Weaknesses: {weaknesses}")
        print(f"Insights: {insights}")
        print(f"Suggestions: {suggestions}")
    else:
        print("Invalid file format. Please upload a PDF or DOCX file.")

if __name__ == '__main__':
    parser = argparse.ArgumentParser(description="Resume Analysis Tool")
    parser.add_argument('file_path', type=str, help='Path to the resume file')
    args = parser.parse_args()
    main(args.file_path)
