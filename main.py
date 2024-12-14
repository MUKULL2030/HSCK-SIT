from flask import Flask, request, jsonify, send_from_directory
import PyPDF2
import docx
import nltk
from nltk.corpus import stopwords
from nltk.tokenize import word_tokenize
import os

app = Flask('__name__', static_folder='static')

nltk.download('punkt')
nltk.download('stopwords')

def validate_file_format(file_path):
    if file_path.endswith('.pdf'):
        try:
            with open(file_path, 'rb') as f:
                PyPDF2.PdfFileReader(f)
            return True
        except PyPDF2.PdfReadError:
            return False
    elif file_path.endswith('.docx'):
        try:
            docx.Document(file_path)
            return True
        except docx.exceptions.DocxPackageNotFoundError:
            return False
    else:
        return False

def extract_text_from_pdf(file_path):
    with open(file_path, 'rb') as f:
        pdf_reader = PyPDF2.PdfFileReader(f)
        text = ''
        for page in range(pdf_reader.numPages):
            text += pdf_reader.getPage(page).extractText()
        return text

def extract_text_from_docx(file_path):
    doc = docx.Document(file_path)
    text = ''
    for para in doc.paragraphs:
        text += para.text
    return text

def extract_keywords(text):
    stop_words = set(nltk.corpus.stopwords.words('english'))
    word_tokens = nltk.word_tokenize(text)
    keywords = [word for word in word_tokens if word not in stop_words]
    return keywords

def get_suggestions(keywords, prompt):
    return prompt + ' ' + ' '.join(keywords)

@app.route('/')
def index():
    return send_from_directory(app.static_folder, 'main.html')

@app.route('/analyze', methods=['POST'])
def analyze_file():
    try:
        file = request.files['file']
        file_path = os.path.join('uploads', file.filename)
        file.save(file_path)
        if validate_file_format(file_path):
            if file_path.endswith('.pdf'):
                text = extract_text_from_pdf(file_path)
            elif file_path.endswith('.docx'):
                text = extract_text_from_docx(file_path)
            keywords = extract_keywords(text)
            strengths = get_suggestions(keywords, 'Strengths:')
            weaknesses = get_suggestions(keywords, 'Weaknesses:')
            betterment_suggestions = get_suggestions(keywords, 'Betterment suggestions:')
            networks = get_suggestions(keywords, 'Networks:')
            courses = get_suggestions(keywords, 'Courses:')
            certifications = get_suggestions(keywords, 'Certifications:')
            return jsonify({
                'strengths': strengths,
                'weaknesses': weaknesses,
                'betterment_suggestions': betterment_suggestions,
                'networks': networks,
                'courses': courses,
                'certifications': certifications
            })
        else:
            return jsonify({'error': 'Invalid file format'}), 400
    except Exception as e:
        print(f"Error: {e}")
        return jsonify({'error': 'An error occurred while processing the file.'}), 500

if __name__ == '_main_':
    app.run(debug=True)
