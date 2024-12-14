from flask import Flask, request, jsonify, send_from_directory
import PyPDF2
import docx
import nltk
from nltk.corpus import stopwords
from nltk.tokenize import word_tokenize
import os
import torch
from transformers import BertTokenizer, BertModel, pipeline

app = Flask('__name__', static_folder='static')

nltk.download('punkt')
nltk.download('stopwords')

# Load BERT model and tokenizer
tokenizer = BertTokenizer.from_pretrained('bert-base-uncased')
model = BertModel.from_pretrained('bert-base-uncased')

# Load zero-shot classification model
classifier = pipeline("zero-shot-classification", model="facebook/bart-large-mnli")

def validate_file_format(file_path):
    if file_path.endswith('.pdf'):
        try:
            with open(file_path, 'rb') as f:
                PyPDF2.PdfReader(f)
            return True
        except PyPDF2.errors.PdfReadError:
            return False
    elif file_path.endswith('.docx'):
        try:
            docx.Document(file_path)
            return True
        except docx.exceptions.DocxPackageNotFoundError:
            return False
    else:
        return False

def extract_text_from_pdf(file_path):
    with open(file_path, 'rb') as f:
        pdf_reader = PyPDF2.PdfReader(f)
        text = ''
        for page in range(len(pdf_reader.pages)):
            text += pdf_reader.pages[page].extract_text()
        return text

def extract_text_from_docx(file_path):
    doc = docx.Document(file_path)
    text = ''
    for para in doc.paragraphs:
        text += para.text
    return text

def extract_keywords(text):
    stop_words = set(nltk.corpus.stopwords.words('english'))
    word_tokens = nltk.word_tokenize(text)
    keywords = [word for word in word_tokens if word not in stop_words]
    return keywords

def parse_resume(text):
    keywords = extract_keywords(text)
    
    # Placeholder for a more sophisticated extraction logic
    skills = [word for word in keywords if word.lower() in ['python', 'machine learning', 'nlp', 'data analysis']]
    experience_years = 0
    if 'years' in keywords:
        years_index = keywords.index('years')
        if years_index > 0 and keywords[years_index - 1].isdigit():
            experience_years = int(keywords[years_index - 1])
    projects = ['project' for word in keywords if 'project' in word.lower()]
    
    return skills, experience_years, projects

def generate_insights(skills, experience_years, projects):
    # Combine all information into one text
    text = f"Skills: {', '.join(skills)}. Experience Years: {experience_years}. Projects: {', '.join(projects)}."
    
    # Define the labels (categories) for classification
    labels = ["strengths", "weaknesses", "insights", "suggestions"]
    
    # Get the classification results
    classification_result = classifier(text, candidate_labels=labels)
    
    # Extract the top results for each label
    strengths = classification_result['labels'][0] if 'strengths' in classification_result['labels'] else "No strengths identified"
    weaknesses = classification_result['labels'][1] if 'weaknesses' in classification_result['labels'] else "No weaknesses identified"
    insights = classification_result['labels'][2] if 'insights' in classification_result['labels'] else "No insights identified"
    suggestions = classification_result['labels'][3] if 'suggestions' in classification_result['labels'] else "No suggestions identified"
    
    return strengths, weaknesses, insights, suggestions

@app.route('/', methods=['GET'])
def index():
    return send_from_directory(app.static_folder, 'main.html')

@app.route('/analyze', methods=['POST'])
def analyze_file():
    try:
        file = request.files['file']
        file_path = os.path.join('uploads', file.filename)
        file.save(file_path)
        if validate_file_format(file_path):
            if file_path.endswith('.pdf'):
                text = extract_text_from_pdf(file_path)
            elif file_path.endswith('.docx'):
                text = extract_text_from_docx(file_path)

            # Extracting skills, experience_years, and projects
            skills, experience_years, projects = parse_resume(text)

            # Generate insights using BERT
            strengths, weaknesses, insights, suggestions = generate_insights(skills, experience_years, projects)

            # Creating a JSON object for the extracted data and insights
            response_data = {
                'extracted_data': {
                    'skills': skills,
                    'experience_years': experience_years,
                    'projects': projects
                },
                'insights': {
                    'strengths': strengths,
                    'weaknesses': weaknesses,
                    'insights': insights,
                    'suggestions': suggestions
                }
            }

            return jsonify(response_data)
        else:
            return jsonify({'error': 'Invalid file format'}), 400
    except Exception as e:
        print(f"Error: {e}")
        return jsonify({'error': 'An error occurred while processing the file.'}), 500

if __name__ == '__main__':
    app.run(debug=True)
